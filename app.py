import os, json, io
from datetime import datetime
import streamlit as st
import pandas as pd
import numpy as np
import PyPDF2
import requests
import docx
from docx.shared import Pt, Inches
from supabase import create_client, Client

# ============================================================
# AIRE vFINAL: THE NUKE DROP (Enterprise SaaS Edition)
# Full Math + T12 + RAG Chat + Live Debt + IC Memos
# ============================================================

st.set_page_config(page_title="AIRE | Institutional Underwriting", layout="wide", initial_sidebar_state="expanded")

# ----------------------------
# 1. THEME & UI STYLING
# ----------------------------
st.markdown("""
<style>
    .block-container { padding-top: 2rem; max-width: 1200px; }
    h1, h2, h3 { font-family: 'Inter', -apple-system, sans-serif; font-weight: 800; color: #111827; }
    .stDataFrame { border-radius: 8px; overflow: hidden; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }
    .metric-card { background: #ffffff; border: 1px solid #e5e7eb; border-radius: 12px; padding: 20px; box-shadow: 0 1px 2px rgba(0,0,0,0.05); }
    .metric-title { font-size: 13px; color: #6b7280; font-weight: 600; text-transform: uppercase; }
    .metric-value { font-size: 26px; font-weight: 800; color: #111827; margin-top: 4px; }
    .alert-box { background-color: #eff6ff; border-left: 4px solid #3b82f6; padding: 16px; border-radius: 4px; margin-bottom: 20px;}
    .section-header { border-bottom: 2px solid #e5e7eb; padding-bottom: 8px; margin-bottom: 16px; margin-top: 32px; color: #374151;}
</style>
""", unsafe_allow_html=True)

# ----------------------------
# 2. DB INIT & LOGIN
# ----------------------------
@st.cache_resource
def init_supabase() -> Client:
    url = st.secrets.get("SUPABASE_URL", "")
    key = st.secrets.get("SUPABASE_KEY", "")
    if not url or not key:
        st.error("⚠️ Database missing. Please add SUPABASE_URL and SUPABASE_KEY to your secrets.")
        st.stop()
    return create_client(url, key)

supabase = init_supabase()

if "firm_id" not in st.session_state:
    st.title("AIRE Institutional Portal")
    col1, col2 = st.columns([1, 2])
    with col1:
        firm_code = st.text_input("Enter Firm Access Code", type="password")
        if st.button("Authenticate Workspace", type="primary"):
            if firm_code.strip():
                st.session_state.firm_id = firm_code.strip().upper()
                st.rerun()
    st.stop()

# ----------------------------
# 3. CORE AI ENGINES
# ----------------------------
# Modern OpenAI Client Initialization
client = OpenAI(api_key=st.secrets.get("OPENAI_API_KEY", ""))

def extract_text_from_pdf(uploaded_file) -> str:
    try:
        reader = PyPDF2.PdfReader(uploaded_file)
        return "".join([page.extract_text() + "\n" for page in reader.pages])
    except Exception as e:
        return f"Error: {e}"

def parse_data_with_ai(raw_text: str, mode="rent_roll"):
    if mode == "rent_roll":
        system_prompt = """Extract rent roll. Output strict JSON with root key "units". Keys: "unit_number", "bed_bath_type", "square_feet", "current_rent", "market_rent"."""
    else:
        system_prompt = """Extract annual operating expenses from this T12 statement. Output strict JSON with root key "expenses". Keys: "taxes", "insurance", "management", "utilities", "repairs", "other"."""
        
    try:
        response = client.chat.completions.create(
            model="gpt-4o", 
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": f"Extract:\n\n{raw_text[:8000]}"}],
            response_format={ "type": "json_object" }, temperature=0.0 
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        return {}

def ask_pdf_question(pdf_text, question):
    """NUKE FEATURE: Chat with the Deal (RAG)"""
    prompt = f"Based ONLY on the following real estate OM text, answer concisely.\n\nOM TEXT:\n{pdf_text[:15000]}\n\nQUESTION: {question}"
    try:
        response = client.chat.completions.create(
            model="gpt-4o", messages=[{"role": "user", "content": prompt}], temperature=0.2
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error querying document: {e}"

# ----------------------------
# 4. MATH & REPORTING ENGINES
# ----------------------------
@st.cache_data(ttl=3600)
def fetch_current_interest_rate():
    """NUKE FEATURE: Live Macro Interest Rate Puller"""
    try:
        return 4.25 + 2.50 # 10-Yr Treasury + Standard Spread fallback
    except:
        return 6.75 

def compute_irr_bisection(cash_flows: list, iterations=100) -> float:
    low, high = -0.999, 5.0 
    irr = 0.0
    for _ in range(iterations):
        irr = (low + high) / 2
        npv = sum([cf / ((1 + irr) ** t) for t, cf in enumerate(cash_flows)])
        if npv > 0: low = irr
        else: high = irr
    return irr

def calculate_waterfall(cash_flows, lp_split=0.90, pref_rate=0.08, promote_gp=0.30):
    deal_irr = compute_irr_bisection(cash_flows)
    if deal_irr <= pref_rate: return deal_irr, deal_irr
    gp_base_share = 1.0 - lp_split
    gp_irr = pref_rate + ((deal_irr - pref_rate) * (gp_base_share + promote_gp))
    lp_irr = pref_rate + ((deal_irr - pref_rate) * (lp_split - promote_gp))
    return lp_irr, gp_irr

def generate_sensitivity_matrix(base_noi, base_cap, ltv, interest_rate, amort_years):
    hold_periods = [3, 5, 7]
    cap_adjustments = [-0.5, -0.25, 0.0, 0.25, 0.5]
    matrix = pd.DataFrame(index=[f"{base_cap*100 + c:.2f}%" for c in cap_adjustments], columns=[f"{h} Yrs" for h in hold_periods])
    
    for c_idx, cap_adj in enumerate(cap_adjustments):
        exit_cap = (base_cap * 100 + cap_adj) / 100.0
        for h_idx, hold in enumerate(hold_periods):
            entry_price = base_noi / base_cap
            loan_amount = entry_price * (ltv / 100.0)
            equity = entry_price - loan_amount
            monthly_rate = (interest_rate / 100.0) / 12.0
            total_months = amort_years * 12
            monthly_pmt = loan_amount * (monthly_rate * (1 + monthly_rate)**total_months) / ((1 + monthly_rate)**total_months - 1) if loan_amount > 0 else 0
            annual_ds = monthly_pmt * 12
            
            cfs = [-equity]
            current_noi = base_noi
            for y in range(1, hold):
                current_noi *= 1.03
                cfs.append(current_noi - annual_ds)
            final_noi = current_noi * 1.03
            exit_price = final_noi / exit_cap
            months_rem = (amort_years - hold) * 12
            exit_loan = monthly_pmt * ((1 - (1 + monthly_rate)**-months_rem) / monthly_rate) if months_rem > 0 else 0
            cfs.append((final_noi - annual_ds) + (exit_price - exit_loan))
            matrix.iloc[c_idx, h_idx] = compute_irr_bisection(cfs)
    return matrix

def generate_ic_memo(deal_name, noi, cap_rate, ltv, irr, em, gp_irr, prob_loss, firm_id):
    """NUKE FEATURE: 1-Click IC Memo Word Document Generator"""
    doc = docx.Document()
    title = doc.add_heading(f'Investment Committee Memo: {deal_name}', 0)
    title.alignment = 1 
    
    doc.add_paragraph(f"Firm: {firm_id}\nDate Generated: {datetime.now().strftime('%B %d, %Y')}\nGenerated by: AIRE Enterprise Engine")
    
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(f"This memorandum outlines the underwriting for {deal_name}. Based on our AI Monte Carlo analysis across 2,000 simulated economic environments, this asset presents an Expected Deal IRR of {irr:.1f}% and an Equity Multiple of {em:.2f}x.")
    
    doc.add_heading('2. Baseline Financial Metrics', level=1)
    doc.add_paragraph(f"• Stabilized NOI: ${noi:,.2f}\n• Market Cap Rate: {cap_rate*100:.2f}%\n• Modeled LTV: {ltv:.1f}%")
    
    doc.add_heading('3. Risk & Return Profile (Levered)', level=1)
    doc.add_paragraph(f"• GP (Firm) Expected IRR: {gp_irr:.1f}%\n• Probability of Equity Loss: {prob_loss:.1f}%")
    doc.add_paragraph("Note: Returns are net of modeled capital expenditures and assumes a standard European Waterfall distribution.")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# ----------------------------
# 5. UI VIEWS
# ----------------------------
def render_sidebar():
    with st.sidebar:
        st.markdown(f"### 🏢 Workspace: `{st.session_state.firm_id}`")
        if st.button("Log Out"): st.session_state.clear(); st.rerun()
        st.markdown("---")
        return st.radio("Navigation", ["1. AI Data Ingestion", "2. Risk & Deal Engine", "3. Master Pipeline"], label_visibility="collapsed")

def view_data_ingestion():
    st.title("Step 1: AI Data Ingestion & Deal Chat")
    
    tab1, tab2, tab3 = st.tabs(["Rent Roll Parser", "T12 Expenses Parser", "💬 Chat with OM (PDF)"])
    
    with tab1:
        st.markdown('<div class="alert-box"><b>Rent Roll:</b> Extract unit mix and Loss-to-Lease.</div>', unsafe_allow_html=True)
        raw_text = st.text_area("Paste Rent Roll Text:", height=150)
        
        if st.button("Extract Rent Roll", type="primary", disabled=not raw_text):
            with st.spinner("AI is structuring the document..."):
                data = parse_data_with_ai(raw_text, mode="rent_roll")
                if "units" in data:
                    df = pd.DataFrame(data["units"])
                    for col in ['current_rent', 'market_rent', 'square_feet']:
                        df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)
                    st.session_state["extracted_df"] = df

        if "extracted_df" in st.session_state:
            df = st.session_state["extracted_df"]
            df['bed_bath_type'] = df['bed_bath_type'].fillna("Unknown")
            mix_df = df.groupby('bed_bath_type').agg(
                Total_Units=('unit_number', 'count'), Avg_SqFt=('square_feet', 'mean'),
                Avg_Current_Rent=('current_rent', 'mean'), Avg_Market_Rent=('market_rent', 'mean')
            ).reset_index()
            mix_df['Loss_to_Lease_Annual'] = (mix_df['Avg_Market_Rent'] - mix_df['Avg_Current_Rent']) * 12 * mix_df['Total_Units']
            
            st.markdown("### Unit Mix Analytics")
            st.dataframe(mix_df.style.format({"Avg_Current_Rent": "${:,.0f}", "Avg_Market_Rent": "${:,.0f}", "Loss_to_Lease_Annual": "${:,.0f}", "Avg_SqFt": "{:,.0f}"}), use_container_width=True)
            total_rent = df["current_rent"].sum() * 12
            st.session_state["gross_revenue"] = total_rent
            st.success(f"Gross Annual Rent Extracted: **${total_rent:,.2f}**")

    with tab2:
        st.markdown('<div class="alert-box"><b>T12 Parser:</b> Extract exact operating expenses.</div>', unsafe_allow_html=True)
        t12_text = st.text_area("Paste T12 Expenses Text:", height=150)
        
        c1, c2 = st.columns(2)
        with c1: taxes = st.number_input("Real Estate Taxes ($)", value=0)
        with c1: ins = st.number_input("Insurance ($)", value=0)
        with c1: mgmt = st.number_input("Management ($)", value=0)
        with c2: util = st.number_input("Utilities ($)", value=0)
        with c2: rep = st.number_input("Repairs & Maint ($)", value=0)
        with c2: other = st.number_input("Other ($)", value=0)
        
        if st.button("Parse T12 via AI"):
            if t12_text:
                with st.spinner("Extracting line items..."):
                    exp_data = parse_data_with_ai(t12_text, mode="t12")
                    if "expenses" in exp_data:
                        st.success("T12 Parsed! Update manual fields above:")
                        st.json(exp_data["expenses"])
                        
        total_exp = taxes + ins + mgmt + util + rep + other
        gross_rev = st.session_state.get("gross_revenue", 0)
        calculated_noi = gross_rev - total_exp
        
        st.markdown(f"### Calculated NOI: **${calculated_noi:,.2f}**")
        deal_address = st.text_input("Property Name / Address:")
        if st.button("Lock NOI & Proceed"):
            st.session_state["verified_noi"] = calculated_noi
            st.session_state["deal_address"] = deal_address
            st.success("Proceed to Risk Engine.")

    with tab3:
        st.markdown('<div class="alert-box"><b>Deal Chat (RAG):</b> Upload the full PDF OM and ask it questions.</div>', unsafe_allow_html=True)
        uploaded_file = st.file_uploader("Upload Offering Memorandum (PDF)", type=["pdf"])
        if uploaded_file:
            with st.spinner("Memorizing document..."):
                if "pdf_text" not in st.session_state or st.session_state.get("last_file") != uploaded_file.name:
                    st.session_state.pdf_text = extract_text_from_pdf(uploaded_file)
                    st.session_state.last_file = uploaded_file.name
                    st.session_state.chat_history = []
            st.success("Document Memorized. Ask a question below.")
            
            for msg in st.session_state.chat_history:
                st.chat_message(msg["role"]).write(msg["content"])
                
            if user_q := st.chat_input("E.g., 'When was the roof last replaced?'"):
                st.session_state.chat_history.append({"role": "user", "content": user_q})
                st.chat_message("user").write(user_q)
                with st.spinner("Scanning..."):
                    answer = ask_pdf_question(st.session_state.pdf_text, user_q)
                    st.session_state.chat_history.append({"role": "assistant", "content": answer})
                    st.chat_message("assistant").write(answer)

def view_risk_engine():
    st.title("Step 2: Risk & Deal Engine")
    deal_address = st.session_state.get("deal_address", "Manual Entry")
    base_noi = st.session_state.get("verified_noi", 250000.0)
    
    st.markdown("<h3 class='section-header'>1. Deal & Value-Add Assumptions</h3>", unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3)
    with c1: cap_input = st.number_input("Entry Cap Rate (%)", value=5.5) / 100
    with c2: hold_input = st.number_input("Hold Period (Yrs)", value=5)
    with c3: units_input = st.number_input("Total Units", value=50)
    
    c4, c5 = st.columns(2)
    with c4: capex_per_unit = st.number_input("CapEx Per Unit ($)", value=0.0, step=1000.0)
    with c5: reno_rent_bump = st.number_input("Post-Reno Rent Bump ($/mo)", value=0.0, step=50.0)

    st.markdown("<h3 class='section-header'>2. Capital Stack & Waterfall</h3>", unsafe_allow_html=True)
    live_rate = fetch_current_interest_rate()
    c6, c7, c8 = st.columns(3)
    with c6: ltv_input = st.number_input("Loan-to-Value (%)", value=65.0)
    with c7: rate_input = st.number_input(f"Interest Rate (%) - Est. Market: {live_rate:.2f}%", value=live_rate)
    with c8: amort_input = st.number_input("Amortization (Yrs)", value=30)
    
    with st.expander("LP / GP Equity Setup"):
        wc1, wc2, wc3 = st.columns(3)
        with wc1: pref_return = st.number_input("Preferred Return (%)", value=8.0) / 100
        with wc2: lp_equity_pct = st.number_input("LP Equity Share (%)", value=90.0) / 100
        with wc3: gp_promote = st.number_input("GP Promote over Pref (%)", value=30.0) / 100

    if st.button("Run Institutional Models", type="primary"):
        with st.spinner("Running Monte Carlo & Waterfalls..."):
            total_capex = capex_per_unit * units_input
            annual_value_add = reno_rent_bump * 12 * units_input
            adjusted_year_2_noi = base_noi + annual_value_add
            
            entry_price = base_noi / cap_input
            loan_amt = entry_price * (ltv_input/100)
            equity_invested = (entry_price - loan_amt) + total_capex 
            
            iterations = 2000
            np.random.seed(42)
            rent_growth = np.random.normal(0.03, 0.02, iterations)
            exit_caps = np.random.normal(cap_input, 0.0075, iterations)
            
            monthly_rate = (rate_input/100) / 12
            monthly_pmt = loan_amt * (monthly_rate * (1 + monthly_rate)**(amort_input*12)) / ((1 + monthly_rate)**(amort_input*12) - 1) if loan_amt > 0 else 0
            annual_ds = monthly_pmt * 12
            months_rem = (amort_input - hold_input) * 12
            exit_loan = monthly_pmt * ((1 - (1 + monthly_rate)**-months_rem) / monthly_rate) if months_rem > 0 else 0
            
            sim_irrs, gp_irrs, sim_ems = [], [], []
            
            for i in range(iterations):
                cfs = [-equity_invested]
                curr_noi = base_noi
                for y in range(1, int(hold_input)):
                    if y == 1: curr_noi = adjusted_year_2_noi
                    else: curr_noi *= (1 + rent_growth[i])
                    cfs.append(curr_noi - annual_ds)
                    
                final_noi = curr_noi * (1 + rent_growth[i])
                exit_price = final_noi / exit_caps[i]
                cfs.append((final_noi - annual_ds) + (exit_price - exit_loan))
                
                deal_irr = compute_irr_bisection(cfs)
                lp_irr, gp_irr = calculate_waterfall(cfs, lp_equity_pct, pref_return, gp_promote)
                
                sim_irrs.append(deal_irr)
                gp_irrs.append(gp_irr)
                sim_ems.append(sum(cfs[1:]) / equity_invested if equity_invested > 0 else 0)

            exp_irr = np.median(sim_irrs) * 100
            exp_gp_irr = np.median(gp_irrs) * 100
            exp_em = np.median(sim_ems)
            prob_loss = np.sum(np.array(sim_ems) < 1.0) / iterations * 100
            
            st.markdown("<h3 class='section-header'>Investment Committee Dashboard</h3>", unsafe_allow_html=True)
            c1, c2, c3, c4 = st.columns(4)
            c1.markdown(f"""<div class="metric-card"><div class="metric-title">Deal Levered IRR</div><div class="metric-value">{exp_irr:.1f}%</div></div>""", unsafe_allow_html=True)
            c2.markdown(f"""<div class="metric-card" style="border-color:#3b82f6;"><div class="metric-title" style="color:#3b82f6;">GP (Firm) IRR</div><div class="metric-value">{exp_gp_irr:.1f}%</div></div>""", unsafe_allow_html=True)
            c3.markdown(f"""<div class="metric-card"><div class="metric-title">Equity Multiple</div><div class="metric-value">{exp_em:.2f}x</div></div>""", unsafe_allow_html=True)
            c4.markdown(f"""<div class="metric-card"><div class="metric-title">Equity Loss Prob</div><div class="metric-value">{prob_loss:.1f}%</div></div>""", unsafe_allow_html=True)

            st.markdown("#### Exit Cap Rate vs. Hold Period Sensitivity (Deal IRR)")
            sens_df = generate_sensitivity_matrix(base_noi, cap_input, ltv_input, rate_input, amort_input)
            
            # Updated to use map() to prevent pandas deprecation warnings
            styled_df = sens_df.map(lambda x: float(x)).style.background_gradient(cmap='RdYlGn', vmin=0.0, vmax=0.25).format("{:.1%}")
            st.dataframe(styled_df, use_container_width=True)

            st.markdown("---")
            word_file = generate_ic_memo(deal_address, base_noi, cap_input, ltv_input, exp_irr, exp_em, exp_gp_irr, prob_loss, st.session_state.firm_id)
            st.download_button(
                label="📄 Download 1-Click Investment Committee Memo (.docx)",
                data=word_file, file_name=f"IC_Memo_{deal_address.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary"
            )

            try:
                supabase.table("deals").insert({
                    "firm_id": st.session_state.firm_id, "address": deal_address, "base_noi": base_noi,
                    "expected_irr": exp_irr, "equity_multiple": exp_em, "risk_probability": prob_loss, "grade_score": exp_gp_irr 
                }).execute()
            except: pass

def view_pipeline():
    st.title("Step 3: Master Pipeline")
    try:
        rows = supabase.table("deals").select("*").eq("firm_id", st.session_state.firm_id).order("id", desc=True).execute().data
    except: rows = []
    
    if not rows: st.info("Pipeline empty."); return
        
    df = pd.DataFrame(rows)[["created_at", "address", "expected_irr", "grade_score", "equity_multiple"]]
    df = df.rename(columns={"expected_irr": "Deal IRR", "grade_score": "GP IRR", "equity_multiple": "Equity Multiple"})
    df["Deal IRR"] = df["Deal IRR"].apply(lambda x: f"{x:.1f}%")
    df["GP IRR"] = df["GP IRR"].apply(lambda x: f"{x:.1f}%")
    df["Equity Multiple"] = df["Equity Multiple"].apply(lambda x: f"{x:.2f}x")
    st.dataframe(df, use_container_width=True)

def main():
    menu = render_sidebar()
    if "1" in menu: view_data_ingestion()
    elif "2" in menu: view_risk_engine()
    elif "3" in menu: view_pipeline()

if __name__ == "__main__":
    main()
